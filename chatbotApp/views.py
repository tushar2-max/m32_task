from django.shortcuts import render
from django.http import JsonResponse
from dotenv import load_dotenv
import os
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.chat_models import ChatOpenAI
from langchain.docstore.document import Document
from docx import Document as DocxDocument
from django.views.decorators.csrf import csrf_exempt
import json  # Import the json module

# Load environment variables from .env
load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")

def home(request):
    """
    Renders the index.html page.
    """
    return render(request, 'index.html')

@csrf_exempt
def chat(request):
    """
    Handles the chat functionality.
    """
    if request.method == 'POST':
        data = json.loads(request.body.decode('utf-8'))
        user_message = data.get("message")
        folder_path = 'chatbotApp/data'  # Path to your .docx files

        if not user_message:
            return JsonResponse({"error": "User message is required!"}, status=400)

        # Extract text, create vectorstore, and chatbot
        texts = extract_text_from_docx(folder_path)
        vectorstore = create_vectorstore(texts)
        chatbot = create_chatbot(vectorstore)

        # System message for chatbot response
        system_message = '''1. Do not mention any course names, course titles, or headings in the response. 
        2. Provide the answer in a general way without referencing the source.
        3. Provide the answer in UK English.
        4. Keep the answers less complicated, use fewer medical words.
        5. Keep the sentences short and clear.'''

        # Get the response from the chatbot
        response = chatbot({
            "question": user_message,
            "messages": [
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_message}
            ],
            "chat_history": []
        })

        # Return the response
        return JsonResponse({
            "response": response['answer'],
        })

    return JsonResponse({"error": "Invalid request method!"}, status=405)

# Helper Functions
def extract_text_from_docx(folder_path):
    texts = []
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(folder_path, filename)
            doc = DocxDocument(file_path)

            content = ""
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):  # Extract headings
                    content += f"\n{para.text.upper()}:\n"
                elif para.style.name == 'List Paragraph':  # Extract bullets
                    content += f"- {para.text}\n"
                else:  # Regular paragraphs
                    content += f"{para.text}\n"

            texts.append({"filename": filename, "content": content.strip()})
    return texts

def create_vectorstore(texts):
    """
    Create a FAISS vectorstore using OpenAI embeddings.
    """
    embeddings = OpenAIEmbeddings(openai_api_key=API_KEY)

    # Convert extracted text into LangChain Document objects
    documents = [
        Document(page_content=text["content"], metadata={"source": text["filename"]})
        for text in texts
    ]

    # Create FAISS vectorstore
    vectorstore = FAISS.from_documents(documents, embeddings)
    return vectorstore

def create_chatbot(vectorstore):
    """
    Create a LangChain chatbot with FAISS retriever.
    """
    chat_openai = ChatOpenAI(
        model="gpt-4",
        openai_api_key=API_KEY,
        temperature=0.2,
    )

    # Create the ConversationalRetrievalChain
    qa_chain = ConversationalRetrievalChain.from_llm(
        chat_openai,
        retriever=vectorstore.as_retriever(),
    )
    return qa_chain
